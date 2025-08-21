from flask import Blueprint, request, jsonify, send_file
from flask_jwt_extended import jwt_required, get_jwt_identity
from app.models import Company, SearchCursor
from app import db
import json
import uuid
import datetime
import io
import gzip
import base64
import pandas as pd
from docx import Document
from docx.shared import Pt

main_bp = Blueprint('main', __name__, url_prefix='/DataAccess')

# 添加一個簡單的測試路由，不需要認證
@main_bp.route('/test', methods=['GET'])
def test_route():
    """
    簡單的測試路由，確認API是否正常運行
    """
    return jsonify({
        'status': 'success',
        'message': '系統正常運行中',
        'timestamp': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }), 200

# 壓縮JSON數據的輔助函數
def compress_data(data):
    json_str = json.dumps(data)
    compressed = gzip.compress(json_str.encode('utf-8'))
    return base64.b64encode(compressed).decode('utf-8')

@main_bp.route('/CreateCursor', methods=['GET'])
@jwt_required()
def create_cursor():
    """
    根據關鍵字創建搜索游標
    """
    collection = request.args.get('collection')
    keywords = request.args.getlist('keywords')
    
    if not collection or not keywords:
        return jsonify({'error': 'Missing required parameters'}), 400
    
    # 根據關鍵字搜索公司
    query = Company.query
    
    for keyword in keywords:
        # 在多個字段中搜索關鍵字
        query = query.filter(
            (Company.company_name.ilike(f'%{keyword}%')) |
            (Company.business_description.ilike(f'%{keyword}%')) |
            (Company.company_address.ilike(f'%{keyword}%'))
        )
    
    # 獲取結果
    results = query.all()
    result_ids = [str(company.id) for company in results]
    total_count = len(results)
    
    # 創建游標記錄
    cursor_id = str(uuid.uuid4())
    expires_at = datetime.datetime.utcnow() + datetime.timedelta(hours=24)  # 游標24小時後過期
    
    new_cursor = SearchCursor(
        cursor_id=cursor_id,
        keywords=json.dumps(keywords),
        result_ids=json.dumps(result_ids),
        total_count=total_count,
        expires_at=expires_at
    )
    
    db.session.add(new_cursor)
    db.session.commit()
    
    return jsonify({
        'cursorId': cursor_id,
        'totalCount': total_count
    }), 200

@main_bp.route('/GetSummary', methods=['GET'])
@jwt_required()
def get_summary():
    """
    根據游標ID獲取分頁數據摘要
    """
    cursor_id = request.args.get('cursorId')
    page = int(request.args.get('page', 1))
    page_size = int(request.args.get('pageSize', 10))
    remove_cursor = request.args.get('removeCursor', 'false').lower() == 'true'
    
    if not cursor_id:
        return jsonify({'error': 'Missing cursor ID'}), 400
    
    # 查找游標
    cursor = SearchCursor.query.filter_by(cursor_id=cursor_id).first()
    
    if not cursor:
        return jsonify({'error': 'Invalid cursor ID'}), 404
    
    # 解析結果ID
    result_ids = json.loads(cursor.result_ids)
    
    # 計算分頁
    start_idx = (page - 1) * page_size
    end_idx = start_idx + page_size
    page_ids = result_ids[start_idx:end_idx]
    
    # 獲取公司數據
    companies = []
    for id in page_ids:
        company = Company.query.get(id)
        if company:
            companies.append({
                'BusinessNo': company.business_no,
                'CompanyName': company.company_name,
                'CompanyAddress': company.company_address,
                'BusinessDescription': company.business_description,
                'AddToCollector': True
            })
    
    # 如果是最後一頁且需要刪除游標
    if remove_cursor and end_idx >= len(result_ids):
        db.session.delete(cursor)
        db.session.commit()
    
    # 壓縮數據
    compressed_data = compress_data(companies)
    
    return compressed_data, 200, {'Content-Type': 'text/plain'}

@main_bp.route('/FindByBusinessNo/<business_no>', methods=['GET'])
@jwt_required()
def find_by_business_no(business_no):
    """
    根據統一編號查詢公司詳細信息
    """
    company = Company.query.filter_by(business_no=business_no).first()
    
    if not company:
        return jsonify({'error': 'Company not found'}), 404
    
    # 轉換為前端需要的格式
    company_data = company.to_dict()
    
    # 壓縮數據
    compressed_data = compress_data(company_data)
    
    return compressed_data, 200, {'Content-Type': 'text/plain'}

@main_bp.route('/TryFindCompanyBusinessNo', methods=['GET'])
@jwt_required()
def try_find_company():
    """
    嘗試根據公司部分名稱查找可能的企業
    """
    company_part_name = request.args.get('companyPartName')
    
    if not company_part_name:
        return jsonify({'error': 'Missing company name'}), 400
    
    # 模糊搜索公司
    companies = Company.query.filter(Company.company_name.ilike(f'%{company_part_name}%')).limit(10).all()
    
    result = []
    for company in companies:
        result.append({
            'businessNo': company.business_no,
            'companyName': company.company_name,
            'city': company.company_address.split('市')[0] + '市' if '市' in company.company_address else company.company_address
        })
    
    return jsonify(result), 200

@main_bp.route('/DownloadCompanyInfos', methods=['POST'])
@jwt_required()
def download_company_infos():
    """
    下載選定企業的詳細信息(Excel格式)
    """
    data = request.get_json()
    
    if not data or 'businessNos' not in data:
        return jsonify({'error': 'Missing business numbers'}), 400
    
    business_nos = data['businessNos']
    
    # 獲取公司數據
    companies = []
    for business_no in business_nos:
        company = Company.query.filter_by(business_no=business_no).first()
        if company:
            companies.append(company.to_dict())
    
    # 創建Excel文件
    df = pd.DataFrame([{
        '統一編號': c['BusinessNo'],
        '公司名稱': c['CompanyName'],
        '地址': c['CompanyAddress'],
        '簡介': c['BusinessDescription'],
        '資本額': c['CapitalAmount'],
        '員工人數': c['EmployeeCount'],
        '組織別': c['OrganizationType'],
        '產業類別': ', '.join(c['Industrials']),
        '聯絡人': ', '.join(c['Contacts']),
        '電話': ', '.join(c['Telephones']),
        '傳真': ', '.join(c['Faxes']),
        '電子信箱': ', '.join(c['Emails']),
        '企業網站': ', '.join(c['Websites'])
    } for c in companies])
    
    # 創建一個BytesIO對象
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, index=False, sheet_name='企業資訊')
    
    output.seek(0)
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name='company_infos.xlsx'
    )

@main_bp.route('/DownloadCompanyLabels', methods=['POST'])
@jwt_required()
def download_company_labels():
    """
    下載企業名條(Word格式)
    """
    data = request.get_json()
    
    if not data or 'businessNos' not in data:
        return jsonify({'error': 'Missing business numbers'}), 400
    
    business_nos = data['businessNos']
    font_size = data.get('fontSize', 12)
    count_per_page = data.get('countPerPage', 9)
    
    # 獲取公司數據
    companies = []
    for business_no in business_nos:
        company = Company.query.filter_by(business_no=business_no).first()
        if company:
            companies.append(company)
    
    # 創建Word文檔
    doc = Document()
    
    # 設置頁面邊距
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)
    
    # 創建表格
    rows_count = (len(companies) + 2) // 3  # 每行3個，向上取整
    if rows_count == 0:
        rows_count = 1
    
    table = doc.add_table(rows=rows_count, cols=3)
    table.style = 'Table Grid'
    
    # 填充表格
    for i, company in enumerate(companies):
        row_idx = i // 3
        col_idx = i % 3
        
        cell = table.cell(row_idx, col_idx)
        paragraph = cell.paragraphs[0]
        
        # 添加公司信息
        run = paragraph.add_run(company.company_name + '\n')
        run.font.size = Pt(font_size)
        run.font.bold = True
        
        run = paragraph.add_run(f"地址: {company.company_address}\n")
        run.font.size = Pt(font_size - 2)
        
        if company.telephones:
            run = paragraph.add_run(f"電話: {company.telephones[0].number}\n")
            run.font.size = Pt(font_size - 2)
        
        if company.faxes:
            run = paragraph.add_run(f"傳真: {company.faxes[0].number}\n")
            run.font.size = Pt(font_size - 2)
    
    # 創建一個BytesIO對象
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='company_labels.docx'
    )